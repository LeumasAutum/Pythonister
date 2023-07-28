import os
from docx import Document
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from datetime import datetime
from pathlib import Path

def count_words_excluding_first_last_page(docx_path):
    doc = Document(docx_path)

    # Exclude the first and last page
    total_words = 0
    for i, paragraph in enumerate(doc.paragraphs):
        if i == 0 or i == len(doc.paragraphs) - 1:
            continue

        # Count words in each paragraph
        words_in_paragraph = len(paragraph.text.split())
        total_words += words_in_paragraph

    return total_words

def truncate_quantity(quantity):
    return int(quantity)

def process_folder_and_create_excel(folder_path):
    wb = Workbook()
    sheet = wb.active

    # Set the headers for the columns
    sheet['A1'] = 'Date'
    sheet['B1'] = 'Task'
    sheet['C1'] = 'Quantity'
    sheet['D1'] = 'Amount'

    # Get a list of files in the folder with their creation dates
    files_with_dates = []
    for file in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file)
        creation_date = os.path.getctime(file_path)
        files_with_dates.append((file, creation_date))

    # Sort files by their creation date in ascending order
    files_with_dates.sort(key=lambda x: x[1])

    # Initialize row number to 2 (because row 1 contains headers)
    row_num = 2
    total_amount = 0

    for file, creation_date in files_with_dates:
        file_path = os.path.join(folder_path, file)
        file_extension = Path(file_path).suffix.lower()

        # Get the creation date of the file
        formatted_date = datetime.fromtimestamp(creation_date).strftime('%m/%d/%Y')

        if file_extension == '.docx':
            # Count words excluding first and last page for Word documents
            total_words = count_words_excluding_first_last_page(file_path)
            quantity = total_words / 275
            quantity = truncate_quantity(quantity)  # Truncate to whole number
            amount = quantity * 300
        elif file_extension == '.pdf':
            # For PDF files, set default quantity and amount
            quantity = 1
            amount = 500
        else:
            # For other file types, set default quantity and amount
            quantity = 0
            amount = 0

        # Write data to the Excel sheet
        sheet.cell(row=row_num, column=1, value=formatted_date)
        sheet.cell(row=row_num, column=2, value=file)
        sheet.cell(row=row_num, column=3, value=quantity)
        sheet.cell(row=row_num, column=4, value=amount)

        # Accumulate the total amount
        total_amount += amount

        # Move to the next row
        row_num += 1

    # Write the total amount two rows after the last data row
    last_data_row = row_num - 1
    total_row = last_data_row + 2
    sheet.cell(row=total_row, column=4, value=total_amount)

    # Save the Excel workbook
    wb.save('Invoice.xlsx')

if __name__ == '__main__':
    folder_path = './'  # Replace with the actual folder path
    process_folder_and_create_excel(folder_path)
